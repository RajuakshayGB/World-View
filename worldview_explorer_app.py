import streamlit as st
import docx
import re
import json
import math
import numpy as np
import plotly.graph_objects as go
from io import BytesIO

# ==============================================================================
# STREAMLIT CONFIGURATION & CUSTOM THEME (MUSEUM OF IDEAS)
# ==============================================================================
st.set_page_config(
    page_title="The Compass of Human Perspectives",
    page_icon="🧭",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for high-end editorial styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Cinzel:wght@400;700&family=Lora:ital,wght@0,400;0,600;1,400&family=Inter:wght@300;400;600&display=swap');
    
    /* Main typography rules */
    .main .block-container {
        font-family: 'Inter', sans-serif;
    }
    
    h1, h2, h3, .museum-title {
        font-family: 'Cinzel', serif;
        font-weight: 700;
        color: #1E293B;
    }
    
    .serif-text {
        font-family: 'Lora', serif;
        font-size: 1.15rem;
        line-height: 1.7;
        color: #334155;
    }
    
    .bilingual-hindi {
        font-family: 'Lora', serif;
        font-size: 1.05rem;
        font-style: italic;
        color: #475569;
        margin-top: 4px;
        line-height: 1.6;
    }
    
    /* Elegant card container */
    .museum-card {
        background-color: #F8FAFC;
        border: 1px solid #E2E8F0;
        border-radius: 12px;
        padding: 24px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05), 0 2px 4px -1px rgba(0, 0, 0, 0.03);
        transition: all 0.3s ease;
    }
    .museum-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.05), 0 4px 6px -2px rgba(0, 0, 0, 0.03);
        border-color: #CBD5E1;
    }
    
    /* Progress and Category Bar style */
    .category-header {
        font-family: 'Cinzel', serif;
        font-size: 0.9rem;
        letter-spacing: 0.15em;
        text-transform: uppercase;
        color: #64748B;
        margin-bottom: 8px;
    }
    
    /* Landing page hero */
    .hero-container {
        text-align: center;
        padding: 80px 20px;
        background: radial-gradient(circle, #FCFCFC 0%, #F1F5F9 100%);
        border-radius: 24px;
        border: 1px solid #E2E8F0;
        margin-bottom: 40px;
    }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 1. CORE SCORING LOGIC & COORDINATES SCHEMA
# ==============================================================================
# We project options into a 4-dimensional trait space:
# Dim 0: TRANSCENDENCE vs. PHYSICALISM (-1.0 = Pure Materialism/Physicalism, +1.0 = Pure Transcendence/Idealism)
# Dim 1: INDIVIDUALISM vs. COLLECTIVISM (-1.0 = Extreme Individualism, +1.0 = Extreme Collectivism)
# Dim 2: TRADITIONALISM vs. PROGRESSIVISM (-1.0 = Traditionalism, +1.0 = Progressivism)
# Dim 3: RATIONALISM vs. EMPIRICISM (-1.0 = Deductive/A-Priori/Intuitive, +1.0 = Inductive/Scientific/Empirical)

WORLDVIEWS = {
    "Secular Scientific Humanism": {
        "vector": [-1.0, +0.1, +0.8, +1.0],
        "thinkers": ["Carl Sagan", "John Dewey", "Richard Dawkins"],
        "description": "A progressive philosophy based on science, reason, human agency, and ethical responsibility, completely rejecting supernatural claims."
    },
    "Stoicism": {
        "vector": [-0.3, -0.4, -0.1, -0.5],
        "thinkers": ["Marcus Aurelius", "Seneca", "Epictetus"],
        "description": "An ancient Greek and Roman philosophy teaching the development of self-control and fortitude to overcome destructive emotions and align with natural cosmic reason."
    },
    "Advaita Vedanta": {
        "vector": [+1.0, -0.2, -0.7, -0.8],
        "thinkers": ["Adi Shankara", "Ramana Maharshi"],
        "description": "An orthodox school of Hindu philosophy asserting that the individual self (Atman) and ultimate absolute reality (Brahman) are identical and non-dual."
    },
    "Marxism": {
        "vector": [-1.0, +1.0, +0.9, +0.5],
        "thinkers": ["Karl Marx", "Friedrich Engels", "Rosa Luxemburg"],
        "description": "A materialist philosophy and socio-economic analysis of class relations and historical progress through social struggle and collective ownership."
    },
    "Daoism": {
        "vector": [+0.4, -0.3, -0.3, -0.4],
        "thinkers": ["Laozi", "Zhuangzi"],
        "description": "A tradition of Chinese origin that emphasizes living in effortless harmony with the Dao (the natural, spontaneous flow of the cosmos)."
    },
    "Early Buddhism": {
        "vector": [+0.2, -0.1, +0.1, -0.6],
        "thinkers": ["Siddhartha Gautama (The Buddha)", "Nagarjuna"],
        "description": "A non-theistic spiritual path focused on overcoming suffering by understanding impermanence, non-attachment, and the illusion of a permanent self (Anatta)."
    },
    "Christian Theism": {
        "vector": [+0.9, +0.1, -0.6, -0.4],
        "thinkers": ["Thomas Aquinas", "Augustine of Hippo", "C.S. Lewis"],
        "description": "A monotheistic faith based on the life and teachings of Jesus Christ, asserting a transcendent personal Creator and moral savior."
    },
    "Ubuntu": {
        "vector": [+0.2, +0.9, +0.3, -0.2],
        "thinkers": ["Desmond Tutu", "Nelson Mandela"],
        "description": "An African communalist philosophy asserting that personhood is relational, encapsulated in the phrase: 'I am because we are.'"
    },
    "Confucianism": {
        "vector": [-0.1, +0.5, -0.9, -0.5],
        "thinkers": ["Confucius", "Mencius"],
        "description": "An East Asian ethical and philosophical system emphasizing filial piety, social order, ritual propriety, and moral governance."
    },
    "Deep Ecology": {
        "vector": [+0.3, +0.4, +0.2, +0.4],
        "thinkers": ["Arne Naess", "Aldo Leopold"],
        "description": "An environmental philosophy advocating for the inherent moral rights of all living beings and ecosystems, rejecting human-centric exploitation."
    },
    "Transhumanism": {
        "vector": [-0.8, -0.2, +1.0, +0.9],
        "thinkers": ["Nick Bostrom", "Ray Kurzweil", "Max More"],
        "description": "An intellectual movement advocating for the enhancement of human biological, cognitive, and physical capabilities using advanced technology."
    },
    "Existentialism": {
        "vector": [-0.3, -0.8, +0.6, -0.2],
        "thinkers": ["Jean-Paul Sartre", "Albert Camus", "Friedrich Nietzsche"],
        "description": "A modern movement asserting that existence precedes essence; humans are radically free and must author their own meaning and moral value."
    },
    "Classical Liberalism": {
        "vector": [-0.6, -0.9, +0.4, +0.5],
        "thinkers": ["John Locke", "Adam Smith", "John Stuart Mill"],
        "description": "A political and economic philosophy championing individual liberty, private property, limited state governance, and voluntary market cooperation."
    }
}

# ==============================================================================
# 2. DOCUMENT PARSING MODULE
# ==============================================================================
@st.cache_data
def load_and_parse_docx(uploaded_file=None):
    """
    Parses questions and options from the Bilingual Word Document.
    Fallback included if no file is provided.
    """
    if uploaded_file is not None:
        doc = docx.Document(BytesIO(uploaded_file.read()))
    else:
        # Fallback to loading the local modified document from artifacts
        try:
            doc = docx.Document("/workspace/artifacts/the-compass-of-human-perspectives-bilingual-v2.docx")
        except Exception:
            return None, None
            
    sections = []
    questions = []
    
    current_section = None
    current_q = None
    
    section_pattern = re.compile(r"^(\d+)\.\s+(.*)$")
    question_pattern = re.compile(r"^Question\s+(\d+):\s*(.*)$")
    option_pattern = re.compile(r"^\(([A-E])\)\s*(.*)$")
    
    i = 0
    while i < len(doc.paragraphs):
        p_text = doc.paragraphs[i].text.strip()
        if not p_text:
            i += 1
            continue
            
        sect_match = section_pattern.match(p_text)
        if sect_match and "Framework" not in p_text and "Table of Contents" not in p_text:
            section_num = int(sect_match.group(1))
            section_name = sect_match.group(2)
            i += 1
            desc_text = ""
            while i < len(doc.paragraphs):
                desc_text = doc.paragraphs[i].text.strip()
                if desc_text:
                    break
                i += 1
            
            current_section = {
                "number": section_num,
                "name": section_name,
                "description": desc_text
            }
            sections.append(current_section)
            i += 1
            continue
            
        q_match = question_pattern.match(p_text)
        if q_match:
            q_num = int(q_match.group(1))
            q_eng = q_match.group(2)
            
            i += 1
            q_hin = ""
            while i < len(doc.paragraphs):
                q_hin = doc.paragraphs[i].text.strip()
                if q_hin:
                    break
                i += 1
                
            current_q = {
                "number": q_num,
                "section": current_section["name"] if current_section else "",
                "section_num": current_section["number"] if current_section else 0,
                "question_english": q_eng,
                "question_hindi": q_hin,
                "options": {}
            }
            questions.append(current_q)
            i += 1
            continue
            
        opt_match = option_pattern.match(p_text)
        if opt_match and current_q:
            opt_letter = opt_match.group(1)
            opt_eng = opt_match.group(2)
            
            i += 1
            opt_hin = ""
            while i < len(doc.paragraphs):
                opt_hin = doc.paragraphs[i].text.strip()
                if opt_hin:
                    break
                i += 1
                
            current_q["options"][opt_letter] = {
                "english": opt_eng,
                "hindi": opt_hin
            }
            i += 1
            continue
            
        i += 1
        
    return sections, questions

# ==============================================================================
# 3. HEURISTIC COORDINATE AND PROFILE CALCULATOR
# ==============================================================================
def calculate_coordinates(answers, questions):
    """
    Maps answers to 4D coordinate space via precise semantic keyword heuristics.
    """
    # Initial vector at the origin
    user_vector = np.array([0.0, 0.0, 0.0, 0.0])
    
    # Keyword classification matrix
    rules = [
        # Dim 0: Transcendence (+) vs. Physicalism (-)
        ("transhumanism", 0, -0.6), ("christian_theism", 0, +0.8), ("advaita_vedanta", 0, +1.0),
        ("daoism", 0, +0.3), ("deep_ecology", 0, +0.2), ("secular_humanism", 0, -1.0),
        ("marxism", 0, -1.0), ("matter", 0, -0.8), ("god", 0, +0.9), ("brahman", 0, +1.0),
        ("pure consciousness", 0, +1.0), ("soul", 0, +0.8), ("afterlife", 0, +0.9),
        
        # Dim 1: Individualism (-) vs. Collectivism (+)
        ("classical_liberalism", 1, -1.0), ("existentialism", 1, -0.8), ("ubuntu", 1, +0.9),
        ("marxism", 1, +0.9), ("confucianism", 1, +0.4), ("socialist", 1, +0.9),
        ("collective", 1, +0.8), ("individual liberty", 1, -1.0), ("bodily autonomy", 1, -0.9),
        ("family", 1, +0.3), ("communal", 1, +0.8), ("private enterprise", 1, -0.8),
        
        # Dim 2: Traditionalism (-) vs. Progressivism (+)
        ("transhumanism", 2, +1.0), ("confucianism", 2, -1.0), ("christian_theism", 2, -0.6),
        ("secular_humanism", 2, +0.6), ("marxism", 2, +0.7), ("respect for tradition", 2, -0.9),
        ("ancestors", 2, -0.9), ("biotechnology", 2, +0.9), ("heritage", 2, -0.8),
        ("rapid change", 2, +0.8), ("genetic engineering", 2, +0.9),
        
        # Dim 3: Rationalism (-) vs. Empiricism (+)
        ("secular_humanism", 3, +0.9), ("scientific", 3, +1.0), ("empirical", 3, +1.0),
        ("advaita_vedanta", 3, -0.8), ("stoicism", 3, -0.5), ("logic", 3, -0.8),
        ("meditative absorption", 3, -0.9), ("scripture", 3, -0.7), ("intuition", 3, -0.6)
    ]
    
    # Accumulate deltas based on text matches
    questions_dict = {q["number"]: q for q in questions}
    for q_num, selected_letter in answers.items():
        if q_num in questions_dict:
            opt_text = questions_dict[q_num]["options"][selected_letter]["english"].lower()
            for kw, dim, delta in rules:
                if kw in opt_text:
                    user_vector[dim] += delta
                    
    # Normalize coordinates to range [-1.0, +1.0] using hyperbolic tangent to cap bounds
    user_vector = np.tanh(user_vector * 0.15)
    return user_vector

# ==============================================================================
# 4. APP INTERFACE LAYOUT
# ==============================================================================
# Sidebar - Document Ingestion
with st.sidebar:
    st.markdown("### 🏛️ PROJECT ATLAS")
    st.write("Customize your Worldview Compass instance by uploading a bilingual DOCX question sheet.")
    uploaded_file = st.file_uploader("Upload Word Document (.docx)", type=["docx"])
    
    st.write("---")
    st.markdown("### 🎛️ TEST CONTROLS")
    reset_btn = st.button("Reset Session State", type="secondary")
    if reset_btn:
        st.session_state.clear()
        st.rerun()

# Load the questionnaire dataset
sections, questions = load_and_parse_docx(uploaded_file)

if questions is None or len(questions) == 0:
    st.error("❌ Critical Error: Could not locate the bilingual document `the-compass-of-human-perspectives-bilingual-v2.docx`. Please upload it in the sidebar!")
    st.stop()

# Track state variables
if "answers" not in st.session_state:
    st.session_state.answers = {}
if "completed_block" not in st.session_state:
    st.session_state.completed_block = 0
if "active_block" not in st.session_state:
    st.session_state.active_block = 1

# Group questions into 5 core Blocks to avoid pagination fatigue
BLOCKS = {
    1: {"name": "Metaphysics & Mind", "range": range(1, 21)},
    2: {"name": "Epistemology & Truth", "range": range(21, 41)},
    3: {"name": "Spirituality & Ethics", "range": range(41, 61)},
    4: {"name": "Governance & Economics", "range": range(61, 81)},
    5: {"name": "Ecology & Future", "range": range(81, 101)}
}

# LANDING PAGE VIEW
if len(st.session_state.answers) < 100 and "started" not in st.session_state:
    st.markdown("""
    <div class='hero-container'>
        <h1 style='font-size: 3rem; margin-bottom: 10px;'>🧭 The Compass of Human Perspectives</h1>
        <p class='serif-text' style='font-size: 1.4rem; color: #475569;'>“Why do you believe what you believe?”</p>
        <p style='max-width: 800px; margin: 30px auto; font-size: 1.1rem; line-height: 1.6; color: #64748B;'>
            Embark on a non-judgmental, intellectually serious exploration of human thought. 
            Across 100 bilingual questions, discover the structural architecture of your worldview 
            and trace your affinities to major global traditions including Stoicism, Advaita Vedanta, Marxism, and Daoism.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        if st.button("Begin the Odyssey →", type="primary", use_container_width=True):
            st.session_state.started = True
            st.rerun()
            
    st.stop()

# THE QUESTIONNAIRE ODYSSEY VIEW
if len(st.session_state.answers) < 100:
    active_b = st.session_state.active_block
    block_info = BLOCKS[active_b]
    
    # Progress indicator
    progress_pct = len(st.session_state.answers) / 100.0
    st.progress(progress_pct)
    st.subheader(f"Block {active_b} of 5: {block_info['name']}")
    st.caption(f"Progress: {len(st.session_state.answers)} / 100 Questions Answered")
    
    st.write("---")
    
    # Render all questions belonging to this block
    questions_in_block = [q for q in questions if q["number"] in block_info["range"]]
    
    for q in questions_in_block:
        st.markdown(f"""
        <div style='margin-top: 30px; margin-bottom: 10px;'>
            <div class='category-header'>{q['section']} (Question {q['number']})</div>
            <div class='serif-text' style='font-weight: 600; font-size: 1.25rem;'>{q['question_english']}</div>
            <div class='bilingual-hindi'>{q['question_hindi']}</div>
        </div>
        """, unsafe_allow_html=True)
        
        # Collect options into key-value pairs
        options_opts = {}
        for l, data in q["options"].items():
            options_opts[l] = f"({l}) {data['english']} / {data['hindi']}"
            
        current_selection = st.session_state.answers.get(q["number"], None)
        selected_index = list(options_opts.keys()).index(current_selection) if current_selection else None
        
        ans = st.radio(
            f"Select Option for Q{q['number']}",
            options=list(options_opts.keys()),
            format_func=lambda x: options_opts[x],
            index=selected_index,
            key=f"radio_q{q['number']}",
            label_visibility="collapsed"
        )
        if ans:
            st.session_state.answers[q["number"]] = ans
            
    st.write("---")
    
    # Block Navigation controls
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if active_b > 1:
            if st.button("← Previous Block", use_container_width=True):
                st.session_state.active_block -= 1
                st.rerun()
    with col3:
        if active_b < 5:
            # Check if all questions in active block have been answered
            unanswered_in_block = [q["number"] for q in questions_in_block if q["number"] not in st.session_state.answers]
            if len(unanswered_in_block) == 0:
                if st.button("Continue Odyssey →", type="primary", use_container_width=True):
                    st.session_state.active_block += 1
                    st.rerun()
            else:
                st.button("Answer All Questions to Proceed", disabled=True, use_container_width=True)
        else:
            unanswered_all = [i for i in range(1, 101) if i not in st.session_state.answers]
            if len(unanswered_all) == 0:
                if st.button("Reveal My Worldview 🧭", type="primary", use_container_width=True):
                    st.rerun()
            else:
                st.button(f"{len(unanswered_all)} Questions Remaining", disabled=True, use_container_width=True)

# THE PROFILE REVEAL VIEW (THE MIRROR & THE CHALLENGE)
else:
    # 1. Compute Coordinates
    user_coords = calculate_coordinates(st.session_state.answers, questions)
    
    st.markdown("""
    <div style='text-align: center; margin-bottom: 40px;'>
        <h1 style='font-size: 3.5rem;'>🧭 A Worldview Has Emerged</h1>
        <p class='serif-text' style='font-size: 1.5rem; color: #475569;'>“Welcome to your cognitive mirror.”</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Calculate similarities using cosine similarity
    affinities = []
    for school, data in WORLDVIEWS.items():
        v = np.array(data["vector"])
        # Cosine similarity
        denom = (np.linalg.norm(user_coords) * np.linalg.norm(v))
        similarity = np.dot(user_coords, v) / denom if denom > 0 else 0.0
        # Convert to percentage [0%, 100%]
        similarity_pct = max(0.0, float(similarity + 1) / 2.0)
        affinities.append((school, similarity_pct, data["description"], data["thinkers"]))
        
    affinities.sort(key=lambda x: x[1], reverse=True)
    primary_school = affinities[0]
    
    # Columns layout: Left = Vector Chart, Right = Profile Description
    col_chart, col_desc = st.columns([1, 1])
    
    with col_chart:
        st.markdown("### 📊 Worldview Vector Space Map")
        
        # Build radar chart of the coordinates
        labels = [
            "Transcendence vs.<br>Physicalism", 
            "Individualism vs.<br>Collectivism", 
            "Traditionalism vs.<br>Progressivism", 
            "Rationalism vs.<br>Empiricism"
        ]
        
        # To make radar chart non-negative, shift [-1, 1] range to [0, 2]
        user_shifted = [x + 1.0 for x in user_coords]
        primary_shifted = [x + 1.0 for x in WORLDVIEWS[primary_school[0]]["vector"]]
        
        # Plotly Radar Chart
        fig = go.Figure()
        fig.add_trace(go.Scatterpolar(
            r=user_shifted + [user_shifted[0]],
            theta=labels + [labels[0]],
            fill='toself',
            name='My Coordinates',
            fillcolor='rgba(30, 41, 59, 0.2)',
            line=dict(color='#1E293B', width=2)
        ))
        
        fig.add_trace(go.Scatterpolar(
            r=primary_shifted + [primary_shifted[0]],
            theta=labels + [labels[0]],
            fill='toself',
            name=primary_school[0],
            fillcolor='rgba(14, 165, 233, 0.15)',
            line=dict(color='#0EA5E9', width=2, dash='dash')
        ))
        
        fig.update_layout(
            polar=dict(
                radialaxis=dict(visible=True, range=[0, 2], showticklabels=False),
                angularaxis=dict(direction="clockwise", period=4)
            ),
            showlegend=True,
            legend=dict(yanchor="top", y=1.1, xanchor="left", x=0.1),
            margin=dict(t=20, b=20, l=40, r=40),
            height=450
        )
        st.plotly_chart(fig, use_container_width=True)
        
    with col_desc:
        st.markdown("### 🧭 Profile Characterization")
        
        # Dynamic label based on coordinates
        char_labels = []
        char_labels.append("Spiritualist" if user_coords[0] > 0.1 else "Physicalist")
        char_labels.append("Communitarian" if user_coords[1] > 0.1 else "Individualist")
        char_labels.append("Progressive" if user_coords[2] > 0.1 else "Traditionalist")
        char_labels.append("Empiricist" if user_coords[3] > 0.1 else "Rationalist")
        
        profile_str = " • ".join(char_labels)
        
        st.markdown(f"""
        <div style='background-color: #1E293B; color: white; padding: 20px; border-radius: 12px; margin-bottom: 24px; text-align: center;'>
            <div style='font-size: 0.8rem; letter-spacing: 0.1em; text-transform: uppercase; color: #94A3B8;'>Your Archetype Profile</div>
            <div style='font-size: 1.4rem; font-weight: 700; margin-top: 4px;'>{profile_str}</div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
        <div class='serif-text'>
            Your primary philosophical affinity resembles <strong>{primary_school[0]}</strong> with a 
            <strong>{primary_school[1]:.1%} similarity</strong> match. 
            <p style='margin-top: 10px; font-size: 1rem; color: #475569;'>{primary_school[2]}</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"**Key Thinkers in this tradition:** {', '.join(primary_school[3])}")
        
    st.write("---")
    
    # Philosophical Affinities Ranking
    st.markdown("### 🏛️ Philosophical Affinities")
    st.write("Your coordinates compared with major global schools of thought:")
    
    col_aff1, col_aff2 = st.columns(2)
    for index, aff in enumerate(affinities[:6]):
        target_col = col_aff1 if index % 2 == 0 else col_aff2
        with target_col:
            st.markdown(f"""
            <div class='museum-card' style='padding: 18px;'>
                <div style='display: flex; justify-content: space-between; align-items: center;'>
                    <div style='font-weight: 700; font-size: 1.1rem; color: #1E293B;'>{index+1}. {aff[0]}</div>
                    <div style='background-color: #F1F5F9; border-radius: 20px; padding: 4px 12px; font-size: 0.85rem; font-weight: 600;'>{aff[1]:.1%} Match</div>
                </div>
                <div style='font-size: 0.9rem; color: #475569; margin-top: 6px;'>{aff[2]}</div>
            </div>
            """, unsafe_allow_html=True)
            
    st.write("---")
    
    # THE CHALLENGE: Cognitive Tension Analysis
    st.markdown("### ⚡ The Challenge (Cognitive Tensions)")
    st.write("Worldviews are not static mathematical formulas. Tension is the catalyst of self-exploration:")
    
    # Rules to trigger dynamic tensions
    tensions_found = 0
    
    # Tension 1: Mystic-Empiricist
    has_brahman = st.session_state.answers.get(1, 'A') == 'B' or st.session_state.answers.get(3, 'A') == 'C'
    has_strict_evidence = st.session_state.answers.get(8, 'A') == 'A' or st.session_state.answers.get(9, 'A') == 'A'
    if has_brahman and has_strict_evidence:
        tensions_found += 1
        st.warning("⚡ **The Mystical-Empirical Threshold**")
        st.markdown("""
        <div class='serif-text' style='font-size: 1rem; color: #334155; margin-bottom: 20px;'>
            You believe that reality is ultimately comprised of a non-dual cosmic consciousness (Brahman) or that matter is an illusion, 
            yet you also assert that scientific replication and empirical evidence are the sole arbiters of truth. 
            Because consciousness itself is non-quantifiable, this places you at the heart of the "Hard Problem of Consciousness."
        </div>
        """, unsafe_allow_html=True)
        
    # Tension 2: Individual Liberty vs Collective Mandatory Care
    has_liberty_first = st.session_state.answers.get(31, 'A') == 'A' or st.session_state.answers.get(49, 'A') == 'A'
    has_collective_welfare = st.session_state.answers.get(53, 'A') == 'D' or st.session_state.answers.get(76, 'A') == 'C'
    if has_liberty_first and has_collective_welfare:
        tensions_found += 1
        st.warning("⚡ **Individual Freedom vs. Collective Solidarity**")
        st.markdown("""
        <div class='serif-text' style='font-size: 1rem; color: #334155; margin-bottom: 20px;'>
            You strongly support the view that fundamental individual rights are inviolable boundaries that should never be traded away, 
            yet you also support state-coordinated economic planning and collective mandates during crises to protect the community. 
            This represents the classic friction between classical liberalism and communitarian social democracy.
        </div>
        """, unsafe_allow_html=True)
        
    # Tension 3: Technological Acceleration vs. Biocentric Limits
    has_acceleration = st.session_state.answers.get(81, 'A') == 'D' or st.session_state.answers.get(89, 'A') == 'A'
    has_deep_ecology = st.session_state.answers.get(81, 'A') == 'A' or st.session_state.answers.get(82, 'A') == 'A'
    if has_acceleration and has_deep_ecology:
        tensions_found += 1
        st.warning("⚡ **Promethean Ambition vs. Ecological Reciprocity**")
        st.markdown("""
        <div class='serif-text' style='font-size: 1rem; color: #334155; margin-bottom: 20px;'>
            You view nature as a platform to be optimized and transcended using biotechnology and life-extension science, 
            yet you also hold that the planetary biosphere has non-negotiable, intrinsic value that humans should unconditionally respect. 
            Balancing transhumanist Promethean ambition with deep ecological humility represents one of the most critical challenges of our century.
        </div>
        """, unsafe_allow_html=True)
        
    if tensions_found == 0:
        st.info("🟢 No major structural tensions detected! Your worldview displays high internal thematic consistency.")
