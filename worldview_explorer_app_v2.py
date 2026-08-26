import streamlit as st
import docx
import re
import json
import math
import numpy as np
import plotly.graph_objects as go
from io import BytesIO

# ==============================================================================
# STREAMLIT CONFIGURATION & CUSTOM THEME (MUSEUM OF IDEAS)
# ==============================================================================
st.set_page_config(
    page_title="The Compass of Human Perspectives",
    page_icon="🧭",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for high-end editorial styling & option button cards
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Cinzel:wght@400;700&family=Lora:ital,wght=0,400;0,600;1,400&family=Inter:wght@300;400;600&display=swap');
    
    /* Main typography rules */
    .main .block-container {
        font-family: 'Inter', sans-serif;
        max-width: 900px;
        padding-top: 3rem;
        padding-bottom: 5rem;
    }
    
    h1, h2, h3, .museum-title {
        font-family: 'Cinzel', serif;
        font-weight: 700;
        color: #1E293B;
        letter-spacing: 0.02em;
    }
    
    .serif-text {
        font-family: 'Lora', serif;
        font-size: 1.15rem;
        line-height: 1.7;
        color: #334155;
    }
    
    .bilingual-hindi {
        font-family: 'Lora', serif;
        font-size: 1.05rem;
        font-style: italic;
        color: #475569;
        margin-top: 4px;
        line-height: 1.6;
    }
    
    /* Elegant card container */
    .museum-card {
        background-color: #F8FAFC;
        border: 1px solid #E2E8F0;
        border-radius: 12px;
        padding: 24px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.02), 0 2px 4px -1px rgba(0, 0, 0, 0.01);
        transition: all 0.3s ease;
    }
    .museum-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.04), 0 4px 6px -2px rgba(0, 0, 0, 0.02);
        border-color: #CBD5E1;
    }
    
    /* Progress and Category Bar style */
    .category-header {
        font-family: 'Cinzel', serif;
        font-size: 0.95rem;
        letter-spacing: 0.15em;
        text-transform: uppercase;
        color: #0EA5E9;
        margin-bottom: 12px;
        font-weight: 700;
    }
    
    /* Landing page hero */
    .hero-container {
        text-align: center;
        padding: 60px 40px;
        background: radial-gradient(circle, #FCFCFC 0%, #F1F5F9 100%);
        border-radius: 24px;
        border: 1px solid #E2E8F0;
        margin-bottom: 40px;
        box-shadow: inset 0 2px 4px rgba(0,0,0,0.01);
    }
    
    /* Elegantly style Option Buttons as left-aligned choice cards */
    .main div[data-testid="stButton"] button {
        text-align: left !important;
        justify-content: flex-start !important;
        align-items: center !important;
        padding: 20px 24px !important;
        border-radius: 12px !important;
        font-size: 1.05rem !important;
        font-family: 'Lora', serif !important;
        font-weight: 500 !important;
        width: 100% !important;
        white-space: normal !important;
        word-wrap: break-word !important;
        display: flex !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.01) !important;
        border: 1px solid #E2E8F0 !important;
        transition: all 0.2s ease-in-out !important;
        line-height: 1.5 !important;
        margin-bottom: 10px !important;
    }
    
    /* Hover state for choice buttons */
    .main div[data-testid="stButton"] button:hover {
        border-color: #0EA5E9 !important;
        background-color: #F0F9FF !important;
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(14, 165, 233, 0.08) !important;
    }
    
    /* Selected choice button (using primary style in Streamlit) */
    .main div[data-testid="stButton"] button[data-testid="baseButton-primary"] {
        background-color: #1E293B !important;
        color: white !important;
        border-color: #1E293B !important;
        font-weight: 600 !important;
        box-shadow: 0 4px 8px rgba(30, 41, 59, 0.1) !important;
    }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 1. CORE SCORING LOGIC & COORDINATES SCHEMA
# ==============================================================================
WORLDVIEWS = {
    "Secular Scientific Humanism": {
        "vector": [-1.0, +0.1, +0.8, +1.0],
        "thinkers": ["Carl Sagan", "John Dewey", "Richard Dawkins"],
        "description": "A progressive philosophy based on science, reason, human agency, and ethical responsibility, completely rejecting supernatural claims."
    },
    "Stoicism": {
        "vector": [-0.3, -0.4, -0.1, -0.5],
        "thinkers": ["Marcus Aurelius", "Seneca", "Epictetus"],
        "description": "An ancient Greek and Roman philosophy teaching the development of self-control and fortitude to overcome destructive emotions and align with natural cosmic reason."
    },
    "Advaita Vedanta": {
        "vector": [+1.0, -0.2, -0.7, -0.8],
        "thinkers": ["Adi Shankara", "Ramana Maharshi"],
        "description": "An orthodox school of Hindu philosophy asserting that the individual self (Atman) and ultimate absolute reality (Brahman) are identical and non-dual."
    },
    "Marxism": {
        "vector": [-1.0, +1.0, +0.9, +0.5],
        "thinkers": ["Karl Marx", "Friedrich Engels", "Rosa Luxemburg"],
        "description": "A materialist philosophy and socio-economic analysis of class relations and historical progress through social struggle and collective ownership."
    },
    "Daoism": {
        "vector": [+0.4, -0.3, -0.3, -0.4],
        "thinkers": ["Laozi", "Zhuangzi"],
        "description": "A tradition of Chinese origin that emphasizes living in effortless harmony with the Dao (the natural, spontaneous flow of the cosmos)."
    },
    "Early Buddhism": {
        "vector": [+0.2, -0.1, +0.1, -0.6],
        "thinkers": ["Siddhartha Gautama (The Buddha)", "Nagarjuna"],
        "description": "A non-theistic spiritual path focused on overcoming suffering by understanding impermanence, non-attachment, and the illusion of a permanent self (Anatta)."
    },
    "Christian Theism": {
        "vector": [+0.9, +0.1, -0.6, -0.4],
        "thinkers": ["Thomas Aquinas", "Augustine of Hippo", "C.S. Lewis"],
        "description": "A monotheistic faith based on the life and teachings of Jesus Christ, asserting a transcendent personal Creator and moral savior."
    },
    "Ubuntu": {
        "vector": [+0.2, +0.9, +0.3, -0.2],
        "thinkers": ["Desmond Tutu", "Nelson Mandela"],
        "description": "An African communalist philosophy asserting that personhood is relational, encapsulated in the phrase: 'I am because we are.'"
    },
    "Confucianism": {
        "vector": [-0.1, +0.5, -0.9, -0.5],
        "thinkers": ["Confucius", "Mencius"],
        "description": "An East Asian ethical and philosophical system emphasizing filial piety, social order, ritual propriety, and moral governance."
    },
    "Deep Ecology": {
        "vector": [+0.3, +0.4, +0.2, +0.4],
        "thinkers": ["Arne Naess", "Aldo Leopold"],
        "description": "An environmental philosophy advocating for the inherent moral rights of all living beings and ecosystems, rejecting human-centric exploitation."
    },
    "Transhumanism": {
        "vector": [-0.8, -0.2, +1.0, +0.9],
        "thinkers": ["Nick Bostrom", "Ray Kurzweil", "Max More"],
        "description": "An intellectual movement advocating for the enhancement of human biological, cognitive, and physical capabilities using advanced technology."
    },
    "Existentialism": {
        "vector": [-0.3, -0.8, +0.6, -0.2],
        "thinkers": ["Jean-Paul Sartre", "Albert Camus", "Friedrich Nietzsche"],
        "description": "A modern movement asserting that existence precedes essence; humans are radically free and must author their own meaning and moral value."
    },
    "Classical Liberalism": {
        "vector": [-0.6, -0.9, +0.4, +0.5],
        "thinkers": ["John Locke", "Adam Smith", "John Stuart Mill"],
        "description": "A political and economic philosophy championing individual liberty, private property, limited state governance, and voluntary market cooperation."
    }
}

# ==============================================================================
# 2. DOCUMENT PARSING MODULE
# ==============================================================================
@st.cache_data
def load_and_parse_docx(uploaded_file=None):
    """
    Parses questions and options from the Bilingual Word Document.
    Fallback included if no file is provided.
    """
    if uploaded_file is not None:
        doc = docx.Document(BytesIO(uploaded_file.read()))
    else:
        # Fallback to loading the local modified document from artifacts or knowledge
        doc = None
        for path in ["/workspace/artifacts/the-compass-of-human-perspectives-bilingual-v2.docx",
                     "/workspace/knowledge/the-compass-of-human-perspectives-bilingual-v2.docx"]:
            try:
                doc = docx.Document(path)
                break
            except Exception:
                continue
        if doc is None:
            return None, None
            
    sections = []
    questions = []
    
    current_section = None
    current_q = None
    
    section_pattern = re.compile(r"^(\d+)\.\s+(.*)$")
    question_pattern = re.compile(r"^Question\s+(\d+):\s*(.*)$")
    option_pattern = re.compile(r"^\(([A-E])\)\s*(.*)$")
    
    i = 0
    while i < len(doc.paragraphs):
        p_text = doc.paragraphs[i].text.strip()
        if not p_text:
            i += 1
            continue
            
        sect_match = section_pattern.match(p_text)
        if sect_match and "Framework" not in p_text and "Table of Contents" not in p_text:
            section_num = int(sect_match.group(1))
            section_name = sect_match.group(2)
            i += 1
            desc_text = ""
            while i < len(doc.paragraphs):
                desc_text = doc.paragraphs[i].text.strip()
                if desc_text:
                    break
                i += 1
            
            current_section = {
                "number": section_num,
                "name": section_name,
                "description": desc_text
            }
            sections.append(current_section)
            i += 1
            continue
            
        q_match = question_pattern.match(p_text)
        if q_match:
            q_num = int(q_match.group(1))
            q_eng = q_match.group(2)
            
            i += 1
            q_hin = ""
            while i < len(doc.paragraphs):
                q_hin = doc.paragraphs[i].text.strip()
                if q_hin:
                    break
                i += 1
                
            current_q = {
                "number": q_num,
                "section": current_section["name"] if current_section else "",
                "section_num": current_section["number"] if current_section else 0,
                "question_english": q_eng,
                "question_hindi": q_hin,
                "options": {}
            }
            questions.append(current_q)
            i += 1
            continue
            
        opt_match = option_pattern.match(p_text)
        if opt_match and current_q:
            opt_letter = opt_match.group(1)
            opt_eng = opt_match.group(2)
            
            i += 1
            opt_hin = ""
            while i < len(doc.paragraphs):
                opt_hin = doc.paragraphs[i].text.strip()
                if opt_hin:
                    break
                i += 1
                
            current_q["options"][opt_letter] = {
                "english": opt_eng,
                "hindi": opt_hin
            }
            i += 1
            continue
            
        i += 1
        
    return sections, questions

# ==============================================================================
# 3. SCALED HEURISTIC COORDINATE AND PROFILE CALCULATOR
# ==============================================================================
def calculate_coordinates_scaled(answers, questions, test_type):
    """
    Maps answers to 4D coordinate space via precise semantic keyword heuristics,
    scaled properly according to the test length (Quick 25 vs Full 100 questions).
    """
    # Initial vector at the origin
    user_vector = np.array([0.0, 0.0, 0.0, 0.0])
    
    # Keyword classification matrix
    rules = [
        # Dim 0: Transcendence (+) vs. Physicalism (-)
        ("transhumanism", 0, -0.6), ("christian_theism", 0, +0.8), ("advaita_vedanta", 0, +1.0),
        ("daoism", 0, +0.3), ("deep_ecology", 0, +0.2), ("secular_humanism", 0, -1.0),
        ("marxism", 0, -1.0), ("matter", 0, -0.8), ("god", 0, +0.9), ("brahman", 0, +1.0),
        ("pure consciousness", 0, +1.0), ("soul", 0, +0.8), ("afterlife", 0, +0.9),
        
        # Dim 1: Individualism (-) vs. Collectivism (+)
        ("classical_liberalism", 1, -1.0), ("existentialism", 1, -0.8), ("ubuntu", 1, +0.9),
        ("marxism", 1, +0.9), ("confucianism", 1, +0.4), ("socialist", 1, +0.9),
        ("collective", 1, +0.8), ("individual liberty", 1, -1.0), ("bodily autonomy", 1, -0.9),
        ("family", 1, +0.3), ("communal", 1, +0.8), ("private enterprise", 1, -0.8),
        
        # Dim 2: Traditionalism (-) vs. Progressivism (+)
        ("transhumanism", 2, +1.0), ("confucianism", 2, -1.0), ("christian_theism", 2, -0.6),
        ("secular_humanism", 2, +0.6), ("marxism", 2, +0.7), ("respect for tradition", 2, -0.9),
        ("ancestors", 2, -0.9), ("biotechnology", 2, +0.9), ("heritage", 2, -0.8),
        ("rapid change", 2, +0.8), ("genetic engineering", 2, +0.9),
        
        # Dim 3: Rationalism (-) vs. Empiricism (+)
        ("secular_humanism", 3, +0.9), ("scientific", 3, +1.0), ("empirical", 3, +1.0),
        ("advaita_vedanta", 3, -0.8), ("stoicism", 3, -0.5), ("logic", 3, -0.8),
        ("meditative absorption", 3, -0.9), ("scripture", 3, -0.7), ("intuition", 3, -0.6)
    ]
    
    # Accumulate deltas based on text matches
    questions_dict = {q["number"]: q for q in questions}
    for q_num, selected_letter in answers.items():
        if q_num in questions_dict:
            opt_text = questions_dict[q_num]["options"][selected_letter]["english"].lower()
            for kw, dim, delta in rules:
                if kw in opt_text:
                    user_vector[dim] += delta
                    
    # Scale coordinates based on the test type length to maintain identical bounds
    total_expected = 100.0 if test_type == "Full" else 25.0
    scaling_factor = 100.0 / total_expected
    
    # Normalize coordinates to range [-1.0, +1.0] using hyperbolic tangent to cap bounds
    user_vector = np.tanh(user_vector * 0.15 * scaling_factor)
    return user_vector

# ==============================================================================
# 4. BILINGUAL DICTIONARY FOR USER INTERFACE
# ==============================================================================
UI_TEXT = {
    "English": {
        "title": "The Compass of Human Perspectives",
        "subtitle": "Why do you believe what you believe?",
        "tagline": "Embark on a non-judgmental, intellectually serious exploration of human thought. Across 25 or 100 questions, discover the structural architecture of your worldview and trace your affinities to major global traditions including Stoicism, Advaita Vedanta, Marxism, and Daoism.",
        "start_btn": "Begin the Odyssey →",
        "quick_test": "Quick Odyssey (25 Questions)",
        "quick_desc": "A swift but philosophically rigorous 10-minute assessment covering one representative question from each of the 25 core dimensions of human thought.",
        "full_test": "Full Odyssey (100 Questions)",
        "full_desc": "The complete, deep-dive experience spanning all 100 questions for maximum precision and a detailed profile map.",
        "test_type_label": "Select your Odyssey length:",
        "reset_btn": "Reset Session State",
        "progress_label": "Question {current} of {total}",
        "answered_label": "Progress: {answered} / {total} Questions Answered",
        "prev_btn": "← Previous Question",
        "next_btn": "Next Question →",
        "reveal_btn": "Reveal My Worldview 🧭",
        "section_prefix": "Section",
        "result_title": "A Worldview Has Emerged",
        "result_subtitle": "Welcome to your cognitive mirror.",
        "map_title": "📊 Worldview Vector Space Map",
        "char_title": "🧭 Profile Characterization",
        "archetype_label": "Your Archetype Profile",
        "affinity_title": "Your primary philosophical affinity resembles <strong>{school}</strong> with a <strong>{similarity:.1%} similarity</strong> match.",
        "thinkers_label": "Key Thinkers in this tradition:",
        "affinities_label": "🏛️ Philosophical Affinities",
        "affinities_desc": "Your coordinates compared with major global schools of thought:",
        "challenge_title": "⚡ The Challenge (Cognitive Tensions)",
        "challenge_desc": "Worldviews are not static mathematical formulas. Tension is the catalyst of self-exploration:",
        "no_tensions": "🟢 No major structural tensions detected! Your worldview displays high internal thematic consistency.",
        "match_label": "Match"
    },
    "Hindi": {
        "title": "मानव दृष्टिकोण का कम्पास (The Compass of Human Perspectives)",
        "subtitle": "आप जो मानते हैं, क्यों मानते हैं?",
        "tagline": "मानव विचार की एक निष्पक्ष, बौद्धिक रूप से गंभीर खोज पर निकलें। 25 या 100 प्रश्नों के माध्यम से अपने विश्वदृष्टिकोण की संरचनात्मक वास्तुकला की खोज करें और स्टोइसिज्म, अद्वैत वेदांत, मार्क्सवाद और सैन्यवाद (Daoism) सहित प्रमुख वैश्विक परंपराओं के साथ अपनी समानता का पता लगाएं।",
        "start_btn": "यात्रा शुरू करें →",
        "quick_test": "त्वरित यात्रा (25 प्रश्न)",
        "quick_desc": "मानव विचार के 25 प्रमुख आयामों में से एक प्रतिनिधि प्रश्न को कवर करने वाला एक त्वरित लेकिन बौद्धिक रूप से कठोर 10 मिनट का मूल्यांकन।",
        "full_test": "पूर्ण यात्रा (100 प्रश्न)",
        "full_desc": "अधिकतम सटीकता और विस्तृत प्रोफ़ाइल मानचित्र के लिए सभी 100 प्रश्नों को कवर करने वाला पूर्ण, गहन अनुभव।",
        "test_type_label": "अपनी यात्रा की अवधि चुनें:",
        "reset_btn": "सत्र स्थिति रीसेट करें",
        "progress_label": "प्रश्न {current} का {total}",
        "answered_label": "प्रगति: {answered} / {total} प्रश्नों के उत्तर दिए गए",
        "prev_btn": "← पिछला प्रश्न",
        "next_btn": "अगला प्रश्न →",
        "reveal_btn": "मेरा विश्वदृष्टिकोण प्रकट करें 🧭",
        "section_prefix": "अनुभाग",
        "result_title": "एक विश्वदृष्टिकोण का उदय हुआ है",
        "result_subtitle": "आपके संज्ञानात्मक दर्पण में आपका स्वागत है।",
        "map_title": "📊 विश्वदृष्टिकोण वेक्टर स्पेस मैप",
        "char_title": "🧭 प्रोफ़ाइल लक्षण वर्णन",
        "archetype_label": "आपका आर्केटाइप प्रोफ़ाइल",
        "affinity_title": "आपकी प्राथमिक दार्शनिक समानता {similarity:.1%} मैच के साथ <strong>{school}</strong> से मिलती जुलती है।",
        "thinkers_label": "इस परंपरा के प्रमुख विचारक:",
        "affinities_label": "🏛️ दार्शनिक समानताएं",
        "affinities_desc": "विचार के प्रमुख वैश्विक स्कूलों के साथ तुलना में आपके निर्देशांक:",
        "challenge_title": "⚡ चुनौती (संज्ञानात्मक तनाव)",
        "challenge_desc": "विश्वदृष्टिकोण स्थिर गणितीय सूत्र नहीं हैं। तनाव आत्म-अन्वेषण का उत्प्रेरक है:",
        "no_tensions": "🟢 कोई बड़ा संरचनात्मक तनाव नहीं पाया गया! आपका विश्वदृष्टिकोण उच्च आंतरिक विषयगत निरंतरता प्रदर्शित करता है।",
        "match_label": "मैच"
    }
}

# ==============================================================================
# 5. APP INTERFACE LAYOUT & STATE INITIALIZATION
# ==============================================================================
# Track state variables
if "answers" not in st.session_state:
    st.session_state.answers = {}
if "current_question_index" not in st.session_state:
    st.session_state.current_question_index = 0
if "test_type" not in st.session_state:
    st.session_state.test_type = "Quick"
if "language" not in st.session_state:
    st.session_state.language = "English"
if "started" not in st.session_state:
    st.session_state.started = False
if "completed" not in st.session_state:
    st.session_state.completed = False

# Sidebar - Settings & Control
with st.sidebar:
    st.markdown("### 🏛️ PROJECT ATLAS")
    st.write("Customize your Worldview Compass instance by uploading a bilingual DOCX question sheet.")
    uploaded_file = st.file_uploader("Upload Word Document (.docx)", type=["docx"])
    
    st.write("---")
    
    st.markdown("### 🌐 LANGUAGE / भाषा")
    selected_lang = st.radio(
        "Preferred Language / भाषा चुनें:",
        ["English", "Hindi"],
        index=0 if st.session_state.language == "English" else 1,
        key="lang_radio_state"
    )
    if selected_lang != st.session_state.language:
        st.session_state.language = selected_lang
        st.rerun()
        
    st.write("---")
    st.markdown("### 🎛️ TEST CONTROLS")
    reset_btn = st.button("Reset Session State", type="secondary")
    if reset_btn:
        st.session_state.clear()
        st.rerun()

# Load UI translation dictionary
ui = UI_TEXT[st.session_state.language]

# Load the questionnaire dataset
sections, questions = load_and_parse_docx(uploaded_file)

if questions is None or len(questions) == 0:
    st.error("❌ Critical Error: Could not locate the bilingual document `the-compass-of-human-perspectives-bilingual-v2.docx`. Please upload it in the sidebar!")
    st.stop()

# Filter active questions based on selected test type
if st.session_state.test_type == "Quick":
    # 25 questions: one representative question (the first) from each of the 25 dimensions
    quick_nums = [1 + i * 4 for i in range(25)]
    active_questions = [q for q in questions if q["number"] in quick_nums]
else:
    # 100 questions
    active_questions = questions

# ==============================================================================
# LANDING PAGE VIEW
# ==============================================================================
if not st.session_state.started and not st.session_state.completed:
    st.markdown(f"""
    <div class='hero-container'>
        <h1 style='font-size: 3rem; margin-bottom: 10px;'>🧭 {ui['title']}</h1>
        <p class='serif-text' style='font-size: 1.4rem; color: #475569;'>“{ui['subtitle']}”</p>
        <p style='max-width: 800px; margin: 30px auto; font-size: 1.1rem; line-height: 1.6; color: #64748B;'>
            {ui['tagline']}
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown(f"<h3 style='text-align:center;'>⚙️ {ui['test_type_label']}</h3>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown(f"""
        <div class='museum-card' style='min-height: 180px;'>
            <h3 style='margin-top:0; color:#0EA5E9;'>⏱️ {ui['quick_test']}</h3>
            <p style='font-size: 0.95rem; color: #475569;'>{ui['quick_desc']}</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button(f"Start: {ui['quick_test']}", type="primary", use_container_width=True, key="start_quick_btn"):
            st.session_state.test_type = "Quick"
            st.session_state.answers = {}  # Clear for fresh run
            st.session_state.current_question_index = 0
            st.session_state.started = True
            st.rerun()
            
    with col2:
        st.markdown(f"""
        <div class='museum-card' style='min-height: 180px;'>
            <h3 style='margin-top:0; color:#1E293B;'>📖 {ui['full_test']}</h3>
            <p style='font-size: 0.95rem; color: #475569;'>{ui['full_desc']}</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button(f"Start: {ui['full_test']}", type="primary", use_container_width=True, key="start_full_btn"):
            st.session_state.test_type = "Full"
            st.session_state.answers = {}  # Clear for fresh run
            st.session_state.current_question_index = 0
            st.session_state.started = True
            st.rerun()
            
    st.stop()

# ==============================================================================
# THE QUESTIONNAIRE ODYSSEY VIEW (ONE-BY-ONE SLIDESHOW)
# ==============================================================================
if st.session_state.started and not st.session_state.completed:
    idx = st.session_state.current_question_index
    q = active_questions[idx]
    
    # Progress indicator
    total_q = len(active_questions)
    progress_pct = len(st.session_state.answers) / float(total_q)
    st.progress(progress_pct)
    
    # Show navigation status
    progress_text = ui["progress_label"].format(current=idx + 1, total=total_q)
    st.subheader(progress_text)
    
    # Show section category
    st.markdown(f"""
    <div style='margin-top: 10px; margin-bottom: 25px;'>
        <div class='category-header'>{ui['section_prefix']} {q['section_num']}/25: {q['section']}</div>
        <div class='serif-text' style='font-weight: 600; font-size: 1.4rem; line-height: 1.5; color: #1E293B;'>
            {q['question_english'] if st.session_state.language == 'English' else q['question_hindi']}
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.write("---")
    
    # Render option buttons
    current_selection = st.session_state.answers.get(q["number"], None)
    
    for letter, data in q["options"].items():
        opt_text = data["english"] if st.session_state.language == 'English' else data["hindi"]
        is_selected = (current_selection == letter)
        
        btn_label = f"({letter}) {opt_text}"
        if is_selected:
            btn_label = f"✅ {btn_label}"
            button_type = "primary"
        else:
            button_type = "secondary"
            
        if st.button(btn_label, key=f"opt_q{q['number']}_{letter}", use_container_width=True, type=button_type):
            st.session_state.answers[q["number"]] = letter
            # Auto-advance to the next question if possible
            if st.session_state.current_question_index < len(active_questions) - 1:
                st.session_state.current_question_index += 1
            st.rerun()
            
    st.write("---")
    
    # Navigation controls
    col_prev, col_spacer, col_next = st.columns([1.5, 2, 1.5])
    with col_prev:
        if st.session_state.current_question_index > 0:
            if st.button(ui["prev_btn"], use_container_width=True, key="btn_prev_question"):
                st.session_state.current_question_index -= 1
                st.rerun()
    with col_next:
        if st.session_state.current_question_index < len(active_questions) - 1:
            has_answered = (q["number"] in st.session_state.answers)
            if st.button(ui["next_btn"], use_container_width=True, key="btn_next_question", disabled=not has_answered):
                st.session_state.current_question_index += 1
                st.rerun()
        else:
            # We are on the last question of the test
            unanswered_all = [x["number"] for x in active_questions if x["number"] not in st.session_state.answers]
            if len(unanswered_all) == 0:
                if st.button(ui["reveal_btn"], type="primary", use_container_width=True, key="btn_reveal_results"):
                    st.session_state.completed = True
                    st.rerun()
            else:
                st.button(f"{len(unanswered_all)} Questions Remaining", disabled=True, use_container_width=True, key="btn_remaining")

# ==============================================================================
# THE PROFILE REVEAL VIEW (THE MIRROR & THE CHALLENGE)
# ==============================================================================
elif st.session_state.completed:
    # 1. Compute Coordinates with scaling
    user_coords = calculate_coordinates_scaled(st.session_state.answers, questions, st.session_state.test_type)
    
    st.markdown(f"""
    <div style='text-align: center; margin-bottom: 40px;'>
        <h1 style='font-size: 3.2rem;'>🧭 {ui['result_title']}</h1>
        <p class='serif-text' style='font-size: 1.4rem; color: #475569;'>“{ui['result_subtitle']}”</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Calculate similarities using cosine similarity
    affinities = []
    for school, data in WORLDVIEWS.items():
        v = np.array(data["vector"])
        denom = (np.linalg.norm(user_coords) * np.linalg.norm(v))
        similarity = np.dot(user_coords, v) / denom if denom > 0 else 0.0
        similarity_pct = max(0.0, float(similarity + 1) / 2.0)
        affinities.append((school, similarity_pct, data["description"], data["thinkers"]))
        
    affinities.sort(key=lambda x: x[1], reverse=True)
    primary_school = affinities[0]
    
    # Radar Chart & Archetype Description
    col_chart, col_desc = st.columns([1.1, 0.9])
    
    with col_chart:
        st.markdown(f"### {ui['map_title']}")
        
        # Radar labels
        if st.session_state.language == "Hindi":
            labels = [
                "पारलौकिकता बनाम भौतिकवाद<br>(Transcendence vs. Physicalism)", 
                "व्यक्तिवाद बनाम समष्टिवाद<br>(Individualism vs. Collectivism)", 
                "पारंपरिकता बनाम प्रगतिशीलता<br>(Traditionalism vs. Progressivism)", 
                "बुद्धिवाद बनाम अनुभववाद<br>(Rationalism vs. Empiricism)"
            ]
        else:
            labels = [
                "Transcendence vs.<br>Physicalism", 
                "Individualism vs.<br>Collectivism", 
                "Traditionalism vs.<br>Progressivism", 
                "Rationalism vs.<br>Empiricism"
            ]
            
        user_shifted = [x + 1.0 for x in user_coords]
        primary_shifted = [x + 1.0 for x in WORLDVIEWS[primary_school[0]]["vector"]]
        
        fig = go.Figure()
        fig.add_trace(go.Scatterpolar(
            r=user_shifted + [user_shifted[0]],
            theta=labels + [labels[0]],
            fill='toself',
            name='My Coordinates' if st.session_state.language == "English" else 'मेरे निर्देशांक',
            fillcolor='rgba(30, 41, 59, 0.18)',
            line=dict(color='#1E293B', width=2.5)
        ))
        
        fig.add_trace(go.Scatterpolar(
            r=primary_shifted + [primary_shifted[0]],
            theta=labels + [labels[0]],
            fill='toself',
            name=primary_school[0],
            fillcolor='rgba(14, 165, 233, 0.12)',
            line=dict(color='#0EA5E9', width=2, dash='dash')
        ))
        
        fig.update_layout(
            polar=dict(
                radialaxis=dict(visible=True, range=[0, 2], showticklabels=False, gridcolor='#E2E8F0'),
                angularaxis=dict(direction="clockwise", period=4, gridcolor='#E2E8F0')
            ),
            showlegend=True,
            legend=dict(yanchor="top", y=1.15, xanchor="left", x=0.05),
            margin=dict(t=30, b=30, l=40, r=40),
            height=430,
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)'
        )
        st.plotly_chart(fig, use_container_width=True)
        
    with col_desc:
        st.markdown(f"### {ui['char_title']}")
        
        # Archetype label
        char_labels = []
        if st.session_state.language == "Hindi":
            char_labels.append("आध्यात्मिक" if user_coords[0] > 0.1 else "भौतिकवादी")
            char_labels.append("सामुदायिक" if user_coords[1] > 0.1 else "व्यक्तिवादी")
            char_labels.append("प्रगतिशील" if user_coords[2] > 0.1 else "पारंपरिक")
            char_labels.append("अनुभववादी" if user_coords[3] > 0.1 else "बुद्धिवादी")
        else:
            char_labels.append("Spiritualist" if user_coords[0] > 0.1 else "Physicalist")
            char_labels.append("Communitarian" if user_coords[1] > 0.1 else "Individualist")
            char_labels.append("Progressive" if user_coords[2] > 0.1 else "Traditionalist")
            char_labels.append("Empiricist" if user_coords[3] > 0.1 else "Rationalist")
            
        profile_str = " • ".join(char_labels)
        
        st.markdown(f"""
        <div style='background-color: #1E293B; color: white; padding: 22px; border-radius: 12px; margin-bottom: 24px; text-align: center; box-shadow: 0 4px 6px rgba(0,0,0,0.05);'>
            <div style='font-size: 0.8rem; letter-spacing: 0.1em; text-transform: uppercase; color: #94A3B8;'>{ui['archetype_label']}</div>
            <div style='font-size: 1.45rem; font-weight: 700; margin-top: 6px; letter-spacing: 0.02em;'>{profile_str}</div>
        </div>
        """, unsafe_allow_html=True)
        
        affinity_text = ui["affinity_title"].format(school=primary_school[0], similarity=primary_school[1])
        st.markdown(f"""
        <div class='serif-text' style='font-size: 1.1rem; line-height: 1.6;'>
            {affinity_text}
            <p style='margin-top: 12px; font-size: 1rem; color: #475569;'>{primary_school[2]}</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"**{ui['thinkers_label']}** {', '.join(primary_school[3])}")
        
    st.write("---")
    
    # Philosophical Affinities Grid
    st.markdown(f"### {ui['affinities_label']}")
    st.write(ui["affinities_desc"])
    
    col_aff1, col_aff2 = st.columns(2)
    for index, aff in enumerate(affinities[:6]):
        target_col = col_aff1 if index % 2 == 0 else col_aff2
        with target_col:
            st.markdown(f"""
            <div class='museum-card' style='padding: 18px;'>
                <div style='display: flex; justify-content: space-between; align-items: center;'>
                    <div style='font-weight: 700; font-size: 1.1rem; color: #1E293B;'>{index+1}. {aff[0]}</div>
                    <div style='background-color: #F1F5F9; border-radius: 20px; padding: 4px 12px; font-size: 0.85rem; font-weight: 600;'>{aff[1]:.1%} {ui['match_label']}</div>
                </div>
                <div style='font-size: 0.9rem; color: #475569; margin-top: 6px;'>{aff[2]}</div>
            </div>
            """, unsafe_allow_html=True)
            
    st.write("---")
    
    # Cognitive Tensions Analysis
    st.markdown(f"### {ui['challenge_title']}")
    st.write(ui["challenge_desc"])
    
    tensions_found = 0
    
    # Tension 1: Mystic-Empiricist
    has_brahman = st.session_state.answers.get(1) == 'B' or st.session_state.answers.get(3) == 'C'
    has_strict_evidence = st.session_state.answers.get(8) == 'A' or st.session_state.answers.get(9) == 'A'
    if has_brahman and has_strict_evidence:
        tensions_found += 1
        st.warning("⚡ **The Mystical-Empirical Threshold**")
        if st.session_state.language == "Hindi":
            st.markdown("""
            <div class='serif-text' style='font-size: 1rem; color: #334155; margin-bottom: 20px;'>
                आप मानते हैं कि वास्तविकता अंततः एक गैर-द्वैत ब्रह्मांडीय चेतना (ब्रह्म) से बनी है या भौतिक जगत एक भ्रम है, 
                फिर भी आप यह भी दावा करते हैं कि वैज्ञानिक प्रतिकृति और अनुभवजन्य साक्ष्य सत्य के एकमात्र निर्णायक हैं। 
                चूंकि चेतना स्वयं मापने योग्य नहीं है, यह आपको "चेतना की कठिन समस्या" के केंद्र में खड़ा करती है।
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class='serif-text' style='font-size: 1rem; color: #334155; margin-bottom: 20px;'>
                You believe that reality is ultimately comprised of a non-dual cosmic consciousness (Brahman) or that matter is an illusion, 
                yet you also assert that scientific replication and empirical evidence are the sole arbiters of truth. 
                Because consciousness itself is non-quantifiable, this places you at the heart of the "Hard Problem of Consciousness."
            </div>
            """, unsafe_allow_html=True)
            
    # Tension 2: Individual Liberty vs Collective Mandatory Care
    has_liberty_first = st.session_state.answers.get(31) == 'A' or st.session_state.answers.get(49) == 'A'
    has_collective_welfare = st.session_state.answers.get(53) == 'D' or st.session_state.answers.get(76) == 'C'
    if has_liberty_first and has_collective_welfare:
        tensions_found += 1
        st.warning("⚡ **Individual Freedom vs. Collective Solidarity**")
        if st.session_state.language == "Hindi":
            st.markdown("""
            <div class='serif-text' style='font-size: 1rem; color: #334155; margin-bottom: 20px;'>
                आप इस विचार का दृढ़ता से समर्थन करते हैं कि मौलिक व्यक्तिगत अधिकार अनुल्लंघनीय सीमाएं हैं जिनका कभी भी व्यापार नहीं किया जाना चाहिए, 
                फिर भी आप संकट के दौरान समुदाय की सुरक्षा के लिए राज्य-समन्वित आर्थिक नियोजन और सामूहिक शासनादेशों का समर्थन करते हैं। 
                यह शास्त्रीय उदारवाद और सामाजिक लोकतंत्र के बीच क्लासिक घर्षण का प्रतिनिधित्व करता है।
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class='serif-text' style='font-size: 1rem; color: #334155; margin-bottom: 20px;'>
                You strongly support the view that fundamental individual rights are inviolable boundaries that should never be traded away, 
                yet you also support state-coordinated economic planning and collective mandates during crises to protect the community. 
                This represents the classic friction between classical liberalism and communitarian social democracy.
            </div>
            """, unsafe_allow_html=True)
            
    # Tension 3: Technological Acceleration vs. Biocentric Limits
    has_acceleration = st.session_state.answers.get(81) == 'D' or st.session_state.answers.get(89) == 'A'
    has_deep_ecology = st.session_state.answers.get(81) == 'A' or st.session_state.answers.get(82) == 'A'
    if has_acceleration and has_deep_ecology:
        tensions_found += 1
        st.warning("⚡ **Promethean Ambition vs. Ecological Reciprocity**")
        if st.session_state.language == "Hindi":
            st.markdown("""
            <div class='serif-text' style='font-size: 1rem; color: #334155; margin-bottom: 20px;'>
                आप जैव प्रौद्योगिकी और जीवन-विस्तार विज्ञान का उपयोग करके प्रकृति को अनुकूलित और पार करने के मंच के रूप में देखते हैं, 
                फिर भी आप यह भी मानते हैं कि ग्रहों के जीवमंडल का गैर-परक्राम्य, अंतर्निहित मूल्य है जिसका मनुष्यों को बिना शर्त सम्मान करना चाहिए। 
                गहन पारिस्थितिक विनम्रता के साथ ट्रांसह्यूमनिस्ट प्रोमेथियन महत्वाकांक्षा को संतुलित करना हमारी शताब्दी की सबसे महत्वपूर्ण चुनौतियों में से एक है।
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class='serif-text' style='font-size: 1rem; color: #334155; margin-bottom: 20px;'>
                You view nature as a platform to be optimized and transcended using biotechnology and life-extension science, 
                yet you also hold that the planetary biosphere has non-negotiable, intrinsic value that humans should unconditionally respect. 
                Balancing transhumanist Promethean ambition with deep ecological humility represents one of the most critical challenges of our century.
            </div>
            """, unsafe_allow_html=True)
            
    if tensions_found == 0:
        st.info(ui["no_tensions"])
        
    st.write("---")
    
    # Restart Button
    col_l, col_m, col_r = st.columns([1, 1, 1])
    with col_m:
        if st.button("Start a New Odyssey / नई यात्रा शुरू करें", type="primary", use_container_width=True, key="btn_restart_final"):
            st.session_state.clear()
            st.rerun()
